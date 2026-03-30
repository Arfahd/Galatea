"""
Tests for markdown to DOCX rendering.
"""

import pytest
from docx import Document
from docx.shared import Pt

from src.utils.markdown_docx import (
    render_markdown_to_docx,
    _add_formatted_runs,
    _add_heading,
    _add_paragraph,
    HEADING_SIZES,
    DEFAULT_FONT_SIZE,
    CODE_FONT,
)


class TestInlineFormatting:
    """Test inline markdown formatting (bold, italic, code)."""

    def test_bold_double_asterisk(self):
        """**bold** should render as bold."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is **bold** text")

        runs = list(p.runs)
        assert len(runs) == 3
        assert runs[0].text == "This is "
        assert runs[0].bold is None or runs[0].bold is False
        assert runs[1].text == "bold"
        assert runs[1].bold is True
        assert runs[2].text == " text"

    def test_bold_double_underscore(self):
        """__bold__ should render as bold."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is __bold__ text")

        runs = list(p.runs)
        assert any(r.text == "bold" and r.bold for r in runs)

    def test_italic_single_asterisk(self):
        """*italic* should render as italic."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is *italic* text")

        runs = list(p.runs)
        assert any(r.text == "italic" and r.italic for r in runs)

    def test_italic_single_underscore(self):
        """_italic_ should render as italic."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is _italic_ text")

        runs = list(p.runs)
        assert any(r.text == "italic" and r.italic for r in runs)

    def test_bold_italic_triple_asterisk(self):
        """***bold italic*** should render as bold and italic."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is ***bold italic*** text")

        runs = list(p.runs)
        bold_italic_runs = [r for r in runs if r.bold and r.italic]
        assert len(bold_italic_runs) == 1
        assert bold_italic_runs[0].text == "bold italic"

    def test_inline_code(self):
        """`code` should render with monospace font."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "Use `print()` function")

        runs = list(p.runs)
        code_runs = [r for r in runs if r.font.name == CODE_FONT]
        assert len(code_runs) == 1
        assert code_runs[0].text == "print()"

    def test_multiple_formatting_in_one_line(self):
        """Multiple formatting in one line should all render correctly."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is **bold** and *italic* and `code`")

        runs = list(p.runs)
        assert any(r.text == "bold" and r.bold for r in runs)
        assert any(r.text == "italic" and r.italic for r in runs)
        assert any(r.text == "code" and r.font.name == CODE_FONT for r in runs)

    def test_plain_text_only(self):
        """Plain text without markdown should render as-is."""
        doc = Document()
        p = doc.add_paragraph()
        _add_formatted_runs(p, "This is plain text without formatting")

        runs = list(p.runs)
        full_text = "".join(r.text for r in runs)
        assert full_text == "This is plain text without formatting"


class TestHeadings:
    """Test heading rendering."""

    def test_heading_level_1(self):
        """# Heading should render as large bold text."""
        doc = Document()
        render_markdown_to_docx(doc, "# Heading 1")

        p = doc.paragraphs[0]
        assert p.runs[0].bold is True
        assert p.runs[0].font.size == Pt(HEADING_SIZES[1])

    def test_heading_level_2(self):
        """## Heading should render as medium bold text."""
        doc = Document()
        render_markdown_to_docx(doc, "## Heading 2")

        p = doc.paragraphs[0]
        assert p.runs[0].bold is True
        assert p.runs[0].font.size == Pt(HEADING_SIZES[2])

    def test_heading_level_3(self):
        """### Heading should render as small bold text."""
        doc = Document()
        render_markdown_to_docx(doc, "### Heading 3")

        p = doc.paragraphs[0]
        assert p.runs[0].bold is True
        assert p.runs[0].font.size == Pt(HEADING_SIZES[3])

    def test_heading_with_inline_formatting(self):
        """Heading with **bold** should have double bold (still bold)."""
        doc = Document()
        render_markdown_to_docx(doc, "# Heading with **emphasis**")

        p = doc.paragraphs[0]
        # All runs in heading should be bold
        assert all(r.bold for r in p.runs)


class TestLists:
    """Test list rendering."""

    def test_unordered_list_dash(self):
        """- item should render as bullet list."""
        doc = Document()
        render_markdown_to_docx(doc, "- Item 1\n- Item 2\n- Item 3")

        # Should have 3 paragraphs with List Bullet style
        assert len(doc.paragraphs) == 3
        for p in doc.paragraphs:
            assert p.style.name == "List Bullet"

    def test_unordered_list_asterisk(self):
        """* item should render as bullet list."""
        doc = Document()
        render_markdown_to_docx(doc, "* Item 1\n* Item 2")

        assert len(doc.paragraphs) == 2
        for p in doc.paragraphs:
            assert p.style.name == "List Bullet"

    def test_ordered_list(self):
        """1. item should render as numbered list."""
        doc = Document()
        render_markdown_to_docx(doc, "1. First\n2. Second\n3. Third")

        assert len(doc.paragraphs) == 3
        for p in doc.paragraphs:
            assert p.style.name == "List Number"

    def test_list_with_formatting(self):
        """List items with inline formatting should render correctly."""
        doc = Document()
        render_markdown_to_docx(doc, "- **Bold item**\n- *Italic item*")

        p1, p2 = doc.paragraphs
        assert any(r.bold for r in p1.runs)
        assert any(r.italic for r in p2.runs)


class TestCodeBlocks:
    """Test code block rendering."""

    def test_code_block(self):
        """``` code ``` should render with monospace font."""
        doc = Document()
        render_markdown_to_docx(doc, "```\nprint('hello')\nprint('world')\n```")

        # Find the code paragraph
        code_paragraphs = [
            p for p in doc.paragraphs if any(r.font.name == CODE_FONT for r in p.runs)
        ]
        assert len(code_paragraphs) == 1
        assert "print" in code_paragraphs[0].text


class TestHorizontalRule:
    """Test horizontal rule rendering."""

    def test_horizontal_rule_dashes(self):
        """--- should render as horizontal line."""
        doc = Document()
        render_markdown_to_docx(doc, "Above\n\n---\n\nBelow")

        # Should have 3 paragraphs: Above, rule, Below
        assert len(doc.paragraphs) == 3
        # Middle paragraph should contain the rule character
        assert "─" in doc.paragraphs[1].text

    def test_horizontal_rule_asterisks(self):
        """*** should render as horizontal line."""
        doc = Document()
        render_markdown_to_docx(doc, "***")

        assert "─" in doc.paragraphs[0].text


class TestFullDocument:
    """Test full document rendering."""

    def test_mixed_content(self):
        """Document with mixed content should render all elements."""
        content = """# Main Title

This is a paragraph with **bold** and *italic* text.

## Section 1

- Bullet point 1
- Bullet point 2

### Subsection

1. Numbered item
2. Another item

---

Here's some `inline code` and a code block:

```
def hello():
    print("Hello!")
```

The end.
"""
        doc = Document()
        render_markdown_to_docx(doc, content)

        # Should have multiple paragraphs
        assert len(doc.paragraphs) > 5

        # Check that various elements exist
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Main Title" in all_text
        assert "Bullet point" in all_text
        assert "Numbered item" in all_text
        assert "inline code" in all_text
        assert "def hello" in all_text

    def test_empty_content(self):
        """Empty content should not raise error."""
        doc = Document()
        render_markdown_to_docx(doc, "")
        assert len(doc.paragraphs) == 0

    def test_plain_text_fallback(self):
        """Plain text without markdown should still render."""
        doc = Document()
        render_markdown_to_docx(
            doc, "Just plain text\nwithout any\nmarkdown formatting"
        )

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "plain text" in all_text


class TestEdgeCases:
    """Test edge cases and error handling."""

    def test_unclosed_bold(self):
        """Unclosed **bold should render as-is."""
        doc = Document()
        render_markdown_to_docx(doc, "This is **unclosed bold")

        all_text = "".join(p.text for p in doc.paragraphs)
        # Should contain the text (with or without **)
        assert "unclosed" in all_text

    def test_empty_formatting(self):
        """Empty formatting markers should not crash."""
        doc = Document()
        render_markdown_to_docx(doc, "Empty ** ** markers")

        # Should not raise error
        assert len(doc.paragraphs) > 0

    def test_nested_formatting(self):
        """Nested formatting like **bold *and italic*** should handle gracefully."""
        doc = Document()
        # This is tricky - just ensure it doesn't crash
        render_markdown_to_docx(doc, "**bold *nested* bold**")

        assert len(doc.paragraphs) > 0

    def test_consecutive_newlines(self):
        """Multiple consecutive newlines should not create excessive paragraphs."""
        doc = Document()
        render_markdown_to_docx(doc, "Para 1\n\n\n\nPara 2")

        # Should have 2 paragraphs, not 4
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) == 2
