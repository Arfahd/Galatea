"""
Markdown to DOCX renderer.

Converts markdown-formatted text to properly formatted python-docx Document.
Supports: headings, bold, italic, code, lists, horizontal rules.
"""

import re
import logging
from typing import Optional
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

logger = logging.getLogger(__name__)

# Heading font sizes
HEADING_SIZES = {
    1: 24,
    2: 18,
    3: 14,
    4: 12,
    5: 11,
    6: 11,
}

# Default font size for body text
DEFAULT_FONT_SIZE = 11

# Monospace font for code
CODE_FONT = "Courier New"


def render_markdown_to_docx(doc: Document, content: str, original_doc: Optional[Document] = None) -> None:
    """
    Render markdown content to a python-docx Document.

    Supported markdown:
    - # Heading 1, ## Heading 2, ### Heading 3, etc.
    - **bold**, *italic*, ***bold italic***
    - `inline code`
    - ``` code blocks ```
    - - unordered lists, * unordered lists
    - 1. ordered lists
    - --- horizontal rules
    - <<< TABLE_N >>> - marker to copy N-th table from original_doc
    - <<< IMAGE_N >>> - marker to copy N-th image from original_doc

    Args:
        doc: python-docx Document object
        content: Markdown-formatted text
        original_doc: Optional original Document to pull tables/images from
    """
    try:
        _render_content(doc, content, original_doc)
    except Exception as e:
        logger.error(f"Markdown parsing failed: {e}", exc_info=True)
        # Fallback: render as plain text
        for paragraph in content.split("\n\n"):
            if paragraph.strip():
                p = doc.add_paragraph()
                run = p.add_run(paragraph.strip())
                run.font.size = Pt(DEFAULT_FONT_SIZE)


def _render_content(doc: Document, content: str, original_doc: Optional[Document] = None) -> None:
    """
    Internal function to parse and render markdown content.
    """
    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_block_lines = []

    while i < len(lines):
        line = lines[i]
        stripped_line = line.strip()

        # Handle code blocks
        if stripped_line.startswith("```"):
            if in_code_block:
                _add_code_block(doc, "\n".join(code_block_lines))
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        if not stripped_line:
            i += 1
            continue

        # Check for structural markers (more robust: search anywhere in line)
        marker_match = re.search(r"<<<\s*(TABLE|IMAGE)_(\d+)\s*>>>", stripped_line)
        if marker_match and original_doc:
            # If there's text before the marker on the SAME line, render it
            text_before = stripped_line[:marker_match.start()].strip()
            if text_before:
                _add_paragraph(doc, text_before)
            
            m_type = marker_match.group(1)
            m_idx = int(marker_match.group(2))
            
            if m_type == "TABLE":
                _copy_table(doc, original_doc, m_idx)
            elif m_type == "IMAGE":
                _copy_image(doc, original_doc, m_idx)
            
            # If there's text after the marker, we'll process it in next turn
            # or just skip this line if it was just the marker
            text_after = stripped_line[marker_match.end():].strip()
            if text_after:
                # If there's text after, we "inject" it back into lines to be processed
                lines[i] = text_after
                # Don't increment i, so we process the 'text_after' in the next loop
                continue
            
            i += 1
            continue

        # Horizontal rule
        if stripped_line in ("---", "***", "___", "- - -", "* * *"):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped_line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            _add_heading(doc, text, level)
            i += 1
            continue

        # Unordered list item
        if re.match(r"^[\s]*[-*+]\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^[\s]*[-*+]\s+", lines[i]):
                item_text = re.sub(r"^[\s]*[-*+]\s+", "", lines[i])
                list_items.append(item_text)
                i += 1
            _add_unordered_list(doc, list_items)
            continue

        # Ordered list item
        if re.match(r"^[\s]*\d+\.\s+", line):
            list_items = []
            while i < len(lines) and re.match(r"^[\s]*\d+\.\s+", lines[i]):
                item_text = re.sub(r"^[\s]*\d+\.\s+", "", lines[i])
                list_items.append(item_text)
                i += 1
            _add_ordered_list(doc, list_items)
            continue

        para_lines = [line]
        i += 1
        while i < len(lines):
            next_line = lines[i]
            next_stripped = next_line.strip()
            if not next_stripped:
                break
            if (
                next_stripped.startswith("#")
                or next_stripped.startswith("```")
                or next_stripped in ("---", "***", "___", "- - -", "* * *")
                or re.match(r"^[\s]*[-*+]\s+", next_line)
                or re.match(r"^[\s]*\d+\.\s+", next_line)
                or re.search(r"<<<\s*(TABLE|IMAGE)_\d+\s*>>>", next_stripped)
            ):
                break
            para_lines.append(next_line)
            i += 1

        _add_paragraph(doc, " ".join(para_lines))


def _copy_table(doc: Document, original_doc: Document, idx: int) -> None:
    """Copy N-th table from original_doc to doc."""
    try:
        if idx < len(original_doc.tables):
            source_table = original_doc.tables[idx]
            p = doc.add_paragraph()
            new_tbl = deepcopy(source_table._element)
            p._p.addnext(new_tbl)
            p._element.getparent().remove(p._element)
            logger.info(f"Copied table {idx} from original doc")
        else:
            logger.warning(f"Table index {idx} out of range (total tables: {len(original_doc.tables)})")
    except Exception as e:
        logger.error(f"Failed to copy table {idx}: {e}")


def _copy_image(doc: Document, original_doc: Document, idx: int) -> None:
    """Copy N-th image from original_doc to doc."""
    try:
        img_found = 0
        for p in original_doc.paragraphs:
            if 'graphic' in p._element.xml:
                if img_found == idx:
                    new_p = doc.add_paragraph()
                    p_element = deepcopy(p._element)
                    new_p._p.getparent().replace(new_p._p, p_element)
                    logger.info(f"Copied image paragraph {idx} from original doc")
                    return
                img_found += 1
        logger.warning(f"Image index {idx} not found in original doc paragraphs")
    except Exception as e:
        logger.error(f"Failed to copy image {idx}: {e}")


def _add_heading(doc: Document, text: str, level: int) -> None:
    """Add a heading to the document."""
    p = doc.add_paragraph()
    _add_formatted_runs(p, text)
    font_size = HEADING_SIZES.get(level, DEFAULT_FONT_SIZE)
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(font_size)


def _add_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with inline formatting."""
    p = doc.add_paragraph()
    _add_formatted_runs(p, text)


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse inline markdown formatting."""
    pattern = (
        r"(\*\*\*(.+?)\*\*\*|___(.+?)___|"
        r"\*\*(.+?)\*\*|__(.+?)__|"
        r"\*(.+?)\*|_([^_]+)_|"
        r"`([^`]+)`|"
        r"([^*_`]+))"
    )

    text_remaining = text

    while text_remaining:
        match = re.search(pattern, text_remaining)

        if not match:
            if text_remaining:
                run = paragraph.add_run(text_remaining)
                run.font.size = Pt(DEFAULT_FONT_SIZE)
            break

        if match.start() > 0:
            before_text = text_remaining[: match.start()]
            run = paragraph.add_run(before_text)
            run.font.size = Pt(DEFAULT_FONT_SIZE)

        if match.group(2) or match.group(3):
            content = match.group(2) or match.group(3)
            run = paragraph.add_run(content)
            run.bold = True
            run.italic = True
        elif match.group(4) or match.group(5):
            content = match.group(4) or match.group(5)
            run = paragraph.add_run(content)
            run.bold = True
        elif match.group(6) or match.group(7):
            content = match.group(6) or match.group(7)
            run = paragraph.add_run(content)
            run.italic = True
        elif match.group(8):
            content = match.group(8)
            run = paragraph.add_run(content)
            run.font.name = CODE_FONT
        elif match.group(9):
            content = match.group(9)
            run = paragraph.add_run(content)

        run.font.size = Pt(DEFAULT_FONT_SIZE)
        text_remaining = text_remaining[match.end() :]


def _add_code_block(doc: Document, code: str) -> None:
    """Add a code block to the document."""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(10)


def _add_horizontal_rule(doc: Document) -> None:
    """Add a horizontal rule to the document."""
    p = doc.add_paragraph()
    run = p.add_run("─" * 50)
    run.font.size = Pt(DEFAULT_FONT_SIZE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_unordered_list(doc: Document, items: list[str]) -> None:
    """Add an unordered (bullet) list to the document."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _add_formatted_runs(p, item)


def _add_ordered_list(doc: Document, items: list[str]) -> None:
    """Add an ordered (numbered) list to the document."""
    for item in items:
        p = doc.add_paragraph(style="List Number")
        _add_formatted_runs(p, item)
